import os
import bs4
import uuid
import json
import torch
import io
import re
from langchain import hub
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.vectorstores import Chroma
from langchain_core.prompts import ChatPromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import MessagesPlaceholder
from flask import Flask, request, jsonify, render_template
from langchain.docstore.document import Document
from pdf2image import convert_from_path
from transformers import TrOCRProcessor, VisionEncoderDecoderModel
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_google_genai import ChatGoogleGenerativeAI
from werkzeug.utils import secure_filename
from langchain.document_loaders import TextLoader, PyPDFLoader, UnstructuredWordDocumentLoader
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from langchain.chains import create_history_aware_retriever
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.prompts import MessagesPlaceholder
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.chat_history import BaseChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_core.messages import HumanMessage, AIMessage
import warnings
warnings.filterwarnings('ignore')


from dotenv import load_dotenv
import os

load_dotenv()  



app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = './uploads'
app.config['ALLOWED_EXTENSIONS'] = {'txt', 'pdf', 'docx'}


# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Metadata path
metadata_path = './document_metadata.json'



os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_SECRET_KEY")


gemini_embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

model = ChatGoogleGenerativeAI(model="gemini-1.5-pro",convert_system_message_to_human=True)



ocr_processor = TrOCRProcessor.from_pretrained("microsoft/trocr-base-handwritten")
ocr_model = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-base-handwritten")



def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def load_document(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        return TextLoader(file_path).load()
    elif ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def load_pdf(file_path):
    # Try extracting text from PDF (non-image-based PDF)
    try:
        loader = PyPDFLoader(file_path)

        documents = loader.load()

        if not documents:
            # Fallback to OCR if text extraction fails
            return load_pdf_with_ocr(file_path)
        return documents
    except Exception:
        # Fallback to OCR if PyPDFLoader fails to extract text
        return load_pdf_with_ocr(file_path)

def load_pdf_with_ocr(pdf_path):
    # Convert PDF pages to images
    images = convert_from_path(pdf_path)
    
    # Process the first image (for simplicity)
    image = images[0]
    
    # OCR extraction
    inputs = ocr_processor(image, return_tensors="pt")
    pixel_values = inputs.pixel_values

    # Generate text using OCR model
    with torch.no_grad():
        generated_ids = ocr_model.generate(pixel_values)
    
    ocr_text = ocr_processor.decode(generated_ids[0], skip_special_tokens=True)
    
    return [{"text": ocr_text, "metadata": {}}]  # Return OCR-extracted text

def load_docx(file_path):
    doc = Document(file_path)
    doc_text = []

    # Check if the document contains text or images
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            doc_text.append({"text": paragraph.text, "metadata": {}})

    if not doc_text:
        # Fallback to OCR if no text is found
        return load_docx_with_ocr(file_path)

    return doc_text

def load_docx_with_ocr(docx_path):
    doc = Document(docx_path)
    ocr_texts = []

    # Check for images in docx and apply OCR to any found
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            # Extract image
            image_data = rel.target_part.blob
            image = Image.open(io.BytesIO(image_data))
            
            # OCR extraction
            inputs = ocr_processor(image, return_tensors="pt")
            pixel_values = inputs.pixel_values

            # Generate text using OCR model
            with torch.no_grad():
                generated_ids = ocr_model.generate(pixel_values)
            
            ocr_text = ocr_processor.decode(generated_ids[0], skip_special_tokens=True)
            
            # Store OCR text
            ocr_texts.append({"text": ocr_text, "metadata": {}})

    return ocr_texts if ocr_texts else [{"text": "No OCR text found.", "metadata": {}}] 



def save_metadata(document_id, filename, filepath, vector_path):
    metadata = {}
    if os.path.exists(metadata_path):
        with open(metadata_path, 'r') as f:
            metadata = json.load(f)
    metadata[document_id] = {
        'filename': filename,
        'filepath': filepath,
        'vectorstore_path': vector_path
    }
    with open(metadata_path, 'w') as f:
        json.dump(metadata, f, indent=4)

def load_metadata(document_id):
    if os.path.exists(metadata_path):
        with open(metadata_path, 'r') as f:
            metadata = json.load(f)
        return metadata.get(document_id)
    return None

def save_response(status, message, document_id=None, error_details=None):
    response = {
        "status": status,
        "message": message,
    }
    if document_id:
        response["document_id"] = document_id
    if error_details:
        response["error_details"] = error_details
    return jsonify(response)


store = {}

@app.route('/')
def home():
    return render_template('index.html', message=None)


@app.route('/api/embedding', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files.get('document')

        if not file or not allowed_file(file.filename):
            return save_response("error", "Invalid file format.")

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        try:
            document_id = str(uuid.uuid4())


            vectorstore_path = f"./chroma_db6/{document_id}"
            os.makedirs(vectorstore_path, exist_ok=True)

            documents = load_document(filepath)

            

            if not documents:
                return save_response("error", "Failed to embed document.", error_details="Document content is empty.")


            text_splitter = RecursiveCharacterTextSplitter(chunk_size=100, chunk_overlap=10)
            docs = text_splitter.split_documents(documents)
            
            # ✅ Add metadata for filtering later
            for doc in docs:
                doc.metadata["document_id"] = document_id
                doc.metadata["filename"] = filename

            vectordb = Chroma.from_documents(docs, gemini_embeddings, persist_directory=vectorstore_path)
            vectordb.persist()


            save_metadata(document_id, filename, filepath, vectorstore_path)
            return save_response("success", "Document embedded successfully.", document_id=document_id)
        except Exception as e:


            return save_response("error", "Failed to embed document.", error_details="Document content is empty.")

    return render_template('index.html', message=None)



@app.route('/api/query', methods=['GET', 'POST'])
def get_chat():

    try:

        if request.method == 'POST':

            query = request.form.get("query")
            session_id = request.form.get("conversation_id")
            document_id = request.form.get("document_id")
            require_citations = request.form.get("require_citations")


            persist_dir = f"./chroma_db6/{document_id}"
            
            
            vectordb = Chroma(
                persist_directory=persist_dir,
                embedding_function=gemini_embeddings
            )
            
            retriever = vectordb.as_retriever(
                search_kwargs={
                    "filter": {"document_id": document_id},                            #Document id Based filteration
                    "k": 3  # number of documents to retrieve
                }
            )

            system_prompt = (
                "You are an assistant for question-answering tasks. "
                "Use the following pieces of retrieved context to answer the question "
                "If you don't know the answer, say that you don't know."
                "Use three sentences maximum and keep the answer concise."
                "\n\n"
                "{context}"
            )
               
            retriever_prompt = (
                "Given a chat history and the latest user question which might reference context in the chat history,"
                "formulate a standalone question which can be understood without the chat history."
                "Do NOT answer the question, just reformulate it if needed and otherwise return it as is."
            )
            
            contextualize_q_prompt  = ChatPromptTemplate.from_messages(
                [
                    ("system", retriever_prompt),
                    MessagesPlaceholder(variable_name="chat_history"),
                    ("human", "{input}"),
            
            
                ]
            )
            
            history_aware_retriever = create_history_aware_retriever(model,retriever,contextualize_q_prompt)
            
            qa_prompt = ChatPromptTemplate.from_messages(
                [
                    ("system", system_prompt),
                    MessagesPlaceholder("chat_history"),
                    ("human", "{input}"),
                ]
            )
            
            question_answer_chain = create_stuff_documents_chain(model, qa_prompt)
            rag_chain = create_retrieval_chain(history_aware_retriever, question_answer_chain)
            
            
            def get_conversation_history(session_id: str) -> BaseChatMessageHistory:
                if session_id not in store:
                    store[session_id] = ChatMessageHistory()              

                return store[session_id]
            
            
            conversational_rag_chain = RunnableWithMessageHistory(
                rag_chain,
                get_conversation_history,
                input_messages_key="input",
                history_messages_key="chat_history",
                output_messages_key="answer",
            )
            
            
            def get_rag_answer(query, session_id):
                response = conversational_rag_chain.invoke(
                    {"input": query},
                    config={
                        "configurable": {"session_id": session_id}
                    }
                )
                return response

            result = get_rag_answer(query, session_id)

            result_answer = result["answer"]


            
            def citations(context_data):
            
                citations = []
                
                for doc in context_data:
                    content = doc.page_content
                    
                    # Find the citations block using regex
                    match = re.search(r'"citations"\s*:\s*\[(.*?)\]', content, re.DOTALL)
                    if match:
                        # Add brackets to make it a valid JSON array
                        citation_block = '[' + match.group(1) + ']'
                        
                        # Replace any incorrect spacing/formatting in keys
                        citation_block = citation_block.replace('page ', 'page').replace('document_name ', 'document_name')
                        
                        try:
                            parsed = json.loads(citation_block)
                            citations.extend(parsed)
                        except json.JSONDecodeError as e:
                            return []

                
                citations_list = []
                for c in citations:
                    citations_list.append(c)

                return citations_list

            context_data = result["context"]


            citations_list = citations(context_data)

            if require_citations:

                
                result =  {
                    "status": "success",
                    "response": {
                    "citations": citations_list,
                    "answer": result_answer

                        }
                        }

                return jsonify(result)
                
            else:


                result =  {
                    "status": "success",
                    "response": {
                    "answer": result_answer,
                    "citations":[]
                        }
                        }

                return jsonify(result)

    except Exception as e:

        result = {
            "status": "error",
            "message": "Invalid conversation ID. Please start a new session.",
            "error":str(e)
        } 

        return jsonify(result)



@app.route('/chat_history', methods=['GET', 'POST'])
def chat_history():

    if request.method == 'POST':

        session_id = request.form.get("session_id")

        if session_id not in store:
            return {"status": "error", "message": f"No conversation found for conversation ID: {session_id}"}

        conversation = []
        for message in store[session_id].messages:
            sender = "AI" if isinstance(message, AIMessage) else "User"
            conversation.append({
                "sender": sender,
                "message": message.content
            })

        return conversation




if __name__ == '__main__':
    app.run(debug=True)